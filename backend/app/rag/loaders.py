"""Convert locally stored source files into LangChain documents."""

from __future__ import annotations

import re
from abc import ABC, abstractmethod
from pathlib import Path

import fitz
from docx import Document as DocxDocument
from langchain_core.documents import Document as LangChainDocument

from app.core.exceptions import ApplicationError
from app.database.models import Document


class BaseDocumentLoader(ABC):
    """Interface implemented by each supported source-file parser."""

    @abstractmethod
    def load(self, file_path: Path, metadata: dict[str, object]) -> list[LangChainDocument]:
        """Parse a source file into one or more non-empty LangChain documents."""

    @staticmethod
    def _require_content(content: str, file_path: Path) -> str:
        if not content.strip():
            raise ApplicationError(
                code="EMPTY_DOCUMENT_CONTENT",
                message=f"The document contains no extractable text: {file_path.name}.",
                status_code=422,
            )
        return content


class PDFDocumentLoader(BaseDocumentLoader):
    """Extract one LangChain document per PDF page using PyMuPDF."""

    def load(self, file_path: Path, metadata: dict[str, object]) -> list[LangChainDocument]:
        try:
            with fitz.open(file_path) as pdf_document:
                pages = [
                    LangChainDocument(
                        page_content=text,
                        metadata={**metadata, "page_number": page_number},
                    )
                    for page_number, page in enumerate(pdf_document, start=1)
                    if (text := page.get_text()).strip()
                ]
        except ApplicationError:
            raise
        except (fitz.FileDataError, RuntimeError, OSError, ValueError) as exc:
            raise ApplicationError(
                code="DOCUMENT_PARSE_FAILED",
                message=f"The PDF file could not be parsed: {file_path.name}.",
                status_code=422,
            ) from exc

        if not pages:
            self._require_content("", file_path)
        return pages


class DOCXDocumentLoader(BaseDocumentLoader):
    """Extract paragraph text from a DOCX file."""

    def load(self, file_path: Path, metadata: dict[str, object]) -> list[LangChainDocument]:
        try:
            docx_document = DocxDocument(file_path)
            content = "\n".join(paragraph.text for paragraph in docx_document.paragraphs)
        except (OSError, ValueError, KeyError, RuntimeError) as exc:
            raise ApplicationError(
                code="DOCUMENT_PARSE_FAILED",
                message=f"The DOCX file could not be parsed: {file_path.name}.",
                status_code=422,
            ) from exc
        return [
            LangChainDocument(
                page_content=self._require_content(content, file_path), metadata=metadata
            )
        ]


class MarkdownDocumentLoader(BaseDocumentLoader):
    """Split Markdown by headings while retaining the nearest section title."""

    _heading_pattern = re.compile(r"^#{1,6}\s+(.+?)\s*$", re.MULTILINE)

    def load(self, file_path: Path, metadata: dict[str, object]) -> list[LangChainDocument]:
        content = _read_text_file(file_path)
        self._require_content(content, file_path)
        headings = list(self._heading_pattern.finditer(content))
        if not headings:
            return [LangChainDocument(page_content=content, metadata=metadata)]

        documents: list[LangChainDocument] = []
        preamble = content[: headings[0].start()].strip()
        if preamble:
            documents.append(LangChainDocument(page_content=preamble, metadata=metadata))
        for index, heading in enumerate(headings):
            section_end = headings[index + 1].start() if index + 1 < len(headings) else len(content)
            section_content = content[heading.start() : section_end].strip()
            if section_content:
                documents.append(
                    LangChainDocument(
                        page_content=section_content,
                        metadata={**metadata, "section_title": heading.group(1).strip()},
                    )
                )
        if not documents:
            self._require_content("", file_path)
        return documents


class TextDocumentLoader(BaseDocumentLoader):
    """Decode common text encodings into one LangChain document."""

    def load(self, file_path: Path, metadata: dict[str, object]) -> list[LangChainDocument]:
        content = _read_text_file(file_path)
        return [
            LangChainDocument(
                page_content=self._require_content(content, file_path), metadata=metadata
            )
        ]


def load_document(document: Document, *, storage_directory: Path) -> list[LangChainDocument]:
    """Resolve a stored document safely and dispatch it to the matching parser."""
    file_path = _resolve_stored_file(document, storage_directory)
    metadata = {
        "document_id": str(document.id),
        "knowledge_base_id": str(document.knowledge_base_id),
        "original_filename": document.original_name,
        "file_type": document.file_type,
    }
    loader = _loader_for_file_type(document.file_type)
    return loader.load(file_path, metadata)


def _loader_for_file_type(file_type: str) -> BaseDocumentLoader:
    loaders: dict[str, BaseDocumentLoader] = {
        "pdf": PDFDocumentLoader(),
        "docx": DOCXDocumentLoader(),
        "markdown": MarkdownDocumentLoader(),
        "txt": TextDocumentLoader(),
    }
    try:
        return loaders[file_type]
    except KeyError as exc:
        raise ApplicationError(
            code="UNSUPPORTED_DOCUMENT_TYPE",
            message=f"Document type is not supported for parsing: {file_type}.",
            status_code=400,
        ) from exc


def _resolve_stored_file(document: Document, storage_directory: Path) -> Path:
    if not document.storage_path:
        raise ApplicationError(
            code="DOCUMENT_STORAGE_PATH_MISSING",
            message="The document does not have a storage path.",
            status_code=422,
        )
    storage_root = storage_directory.resolve()
    file_path = (storage_root / document.storage_path).resolve()
    if not file_path.is_relative_to(storage_root):
        raise ApplicationError(
            code="INVALID_STORAGE_PATH",
            message="The document storage path is invalid.",
            status_code=400,
        )
    if not file_path.is_file():
        raise ApplicationError(
            code="DOCUMENT_FILE_NOT_FOUND",
            message=f"The stored document file was not found: {document.original_name}.",
            status_code=404,
        )
    return file_path


def _read_text_file(file_path: Path) -> str:
    try:
        data = file_path.read_bytes()
    except OSError as exc:
        raise ApplicationError(
            code="DOCUMENT_FILE_NOT_FOUND",
            message=f"The stored document file was not found: {file_path.name}.",
            status_code=404,
        ) from exc

    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ApplicationError(
        code="DOCUMENT_PARSE_FAILED",
        message=f"The text file could not be decoded: {file_path.name}.",
        status_code=422,
    )
